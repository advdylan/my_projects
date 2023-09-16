import sqlite3
import pandas as pd
from docx import Document
from docx.shared import Cm

from generator_zamowien import generator
def main():
    con = sqlite3.connect("zamowienia.db")
    cur = con.cursor()
    zamowienia = generator(100)

    cur.execute("DELETE FROM zamowienia;")
    insert_row = 'INSERT INTO zamowienia(id, rodzaj_drewna, metry, lesnictwo, lokalizacja) VALUES (?, ? ,? ,? ,?)'
    for i in zamowienia:
        cur.execute(insert_row, i)
        print(i)
    con.commit()
    #con.close()



    lesnictwo = input("Podaj nazwe lesnictwa z ktorego chcesz raport: ")
    print(lesnictwo)
    res = cur.execute('SELECT * FROM zamowienia WHERE lesnictwo=?', (lesnictwo,))
    baza = res.fetchall()
    for i in baza:
        print(i[0])

    document = Document()
    document.add_heading('Raport magazynowy', 0)
    p = document.add_paragraph("Raport składowania materiału")
    headers = [desc[0] for desc in cur.description]
    
    table = document.add_table(rows = 1, cols=len(headers))
    print(len(headers))
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = headers[0]
    hdr_cells[1].text = headers[1]
    hdr_cells[2].text = headers[2]
    hdr_cells[3].text = headers[3]
    hdr_cells[4].text = headers[4]
    for row in baza:
        row_cells = table.add_row().cells
        row_cells[0].text = str(row[0])
        row_cells[1].text = str(row[1])
        row_cells[2].text = str(row[2])
        row_cells[3].text = str(row[3])
        row_cells[4].text = str(row[4])


    document.add_page_break()
    document.save("raport_sprzedazy.docx")

if __name__ == '__main__':
    main()









    #cur.execute("CREATE TABLE zamowienia(id INTEGER PRIMARY KEY AUTOINCREMENT NOT NULL, rodzaj_drewna TEXT NOT NULL, metry INTEGER, lesnictwo TEXT NOT NULL, lokalizacja TEXT NOT NULL)")
    #cur.execute("""
    #INSERT INTO zamowienia VALUES
    #('1', 'sosna', '10', 'Warszawa', 'Las_Poniatowski')
    #""")
    #con.commit()
    """
    cur.execute("DELETE FROM zamowienia;")
    insert_row = 'INSERT INTO zamowienia(id, rodzaj_drewna, metry, lesnictwo, lokalizacja) VALUES (?, ? ,? ,? ,?)'
    for i in zamowienia:
        cur.execute(insert_row, i)
        print(i)
    con.commit()
    con.close()
    """


    """
    cur.execute("SELECT * from zamowienia")
    headers = [desc[0] for desc in cur.description]
    df = pd.DataFrame(baza, columns = headers)
    writer = pd.ExcelWriter('C:\Documents and Settings\cnc\Desktop\Py\project\baza_nadlesnictwo\excel.xls')
    df.to_excel(writer, sheet_name = "Sheet1", index = False)
    writer.close()
    """